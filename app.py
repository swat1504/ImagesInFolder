import streamlit as st
import os
import tempfile
from docx import Document
from docx.shared import Inches

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, Image, Spacer, Paragraph, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

from pypdf import PdfReader, PdfWriter

import zipfile
import re
import pandas as pd
from io import BytesIO

import shutil
import math

st.markdown("""
<style>

/* Make all labels bold */
label, .stMarkdown p {
    font-weight: 700 !important;
    margin-bottom: 4px !important;
}

/* Reduce space ABOVE inputs */
div[data-testid="stTextInput"],
div[data-testid="stFileUploader"],
div[data-testid="stRadio"] {
    margin-top: -10px !important;
}

/* Reduce space BELOW markdown labels */
.stMarkdown {
    margin-bottom: -10px !important;
}

/* Optional: tighten uploader boxes */
section[data-testid="stFileUploaderDropzone"] {
    padding-top: 10px !important;
    padding-bottom: 10px !important;
}

/* Make TAB text larger */
button[data-baseweb="tab"] {
    font-size: 20px !important;
    font-weight: 700 !important;
    padding: 12px 24px !important;
}

/* Increase tab height */
div[data-baseweb="tab-list"] {
    gap: 12px;
}

/* Active tab underline thicker */
button[data-baseweb="tab"][aria-selected="true"] {
    border-bottom: 3px solid #ff4b4b !important;
}

/* Optional: icons inside tabs bigger */
button[data-baseweb="tab"] span {
    font-size: 20px !important;
}

</style>
""", unsafe_allow_html=True)

PAGE_WIDTH, PAGE_HEIGHT = A4

VISIT_TYPES = [
    "RESIDENCE",
    "EMPLOYMENT / BUSINESS",
    "RESI CUM OFFICE",
    "OTHERS"
]

PDF_CATEGORIES = [
    "Main File",
    "Income Tax Return",
    "Form 26AS",
    "Quotation",
    "Shop Act License",
    "Uddyam Certificate",
    "Others",
    "Images Report"
]

# ---------------- UTILS ---------------- #

def safe_name(text):
    return text.replace("/", "_").replace(" ", "_")

# ---------------- TAB 1 PDF ---------------- #

def generate_pdf(visit_images, output_path):

    LEFT = RIGHT = TOP = 40
    BOTTOM = 80

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=LEFT,
        rightMargin=RIGHT,
        topMargin=TOP,
        bottomMargin=BOTTOM
    )

    styles = getSampleStyleSheet()
    elements = []

    usable_width = PAGE_WIDTH - LEFT - RIGHT
    usable_height = PAGE_HEIGHT - TOP - BOTTOM

    img_w = usable_width / 2 - 12
    img_h = usable_height / 2 - 40

    def resize(path):
        img = Image(path)
        scale = min(img_w/img.drawWidth, img_h/img.drawHeight)
        img.drawWidth *= scale
        img.drawHeight *= scale
        return img

    first = True

    for visit, images in visit_images.items():

        if not images:
            continue

        if not first:
            elements.append(PageBreak())

        first = False

        title = visit if visit == "Others" else f"{visit} VISIT"
        elements.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
        elements.append(Spacer(1,15))

        batches = list(range(0,len(images),4))

        for idx,i in enumerate(batches):

            batch = images[i:i+4]
            row1,row2=[],[]

            for j,p in enumerate(batch):
                (row1 if j<2 else row2).append(resize(p))

            while len(row1)<2: row1.append("")
            while len(row2)<2: row2.append("")

            t = Table([row1,row2],colWidths=usable_width/2)
            elements.append(t)

            if idx != len(batches)-1:
                elements.append(Spacer(1,20))

    doc.build(elements)

# ---------------- TAB 1 WORD ---------------- #

def generate_word(visit_images, output_path):

    doc = Document()
    IMG_SIZE = Inches(3)
    first = True

    for visit, images in visit_images.items():

        if not images:
            continue

        if not first:
            doc.add_page_break()

        first = False

        title = visit if visit=="Others" else f"{visit} VISIT"
        doc.add_heading(title,level=1)

        for i in range(0,len(images),4):

            batch = images[i:i+4]
            table = doc.add_table(rows=2,cols=2)

            idx=0
            for r in range(2):
                for c in range(2):
                    if idx<len(batch):
                        table.rows[r].cells[c].paragraphs[0].add_run().add_picture(
                            batch[idx], width=IMG_SIZE
                        )
                    idx+=1

    doc.save(output_path)

# ---------------- TAB 2 PDF MERGE ---------------- #

def merge_pdfs(category_files, output_path):

    writer = PdfWriter()

    for category, files in category_files.items():

        for f in files:

            reader = PdfReader(f, strict=False)

            if category == "Form 26AS":
                writer.add_page(reader.pages[0])
            else:
                for page in reader.pages:
                    writer.add_page(page)

    with open(output_path, "wb") as out:
        writer.write(out)

# ---------------- PDF TEXT ---------------- #

def extract_pdf_text(path):
    reader = PdfReader(path)
    text = ""
    for p in reader.pages:
        try:
            text += p.extract_text()
        except:
            pass
    return text


# ---------------- ITR PARSER ---------------- #

def parse_itr_details(text):
    text_norm = re.sub(r"[ \t]+", " ", text)
    text_norm = text_norm.replace("\xa0", " ")

    # PAN
    pan = ""
    m = re.search(r"\bPAN\s+([A-Z]{5}[0-9]{4}[A-Z])\b", text_norm)
    if m:
        pan = m.group(1).strip()

    # Assessment Year
    ay = ""
    m = re.search(r"Assessment\s*Year\s+(\d{4}-\d{2})", text_norm, re.I)
    if m:
        ay = m.group(1).strip()

    # Acknowledgement Number
    ack = ""
    ack_patterns = [
        r"Acknowledgement Number[:\s]+(\d{12,20})",
        r"e-?Filing Acknowledgement Number[:\s]+(\d{12,20})",
        r"Filed u/s.*?e-?Filing Acknowledgement Number\s+(\d{12,20})",
    ]
    for pat in ack_patterns:
        m = re.search(pat, text_norm, re.I | re.S)
        if m:
            ack = m.group(1).strip()
            break

    # Date of filing
    filing_date = ""
    date_patterns = [
        r"Date of filing\s*:\s*([0-9]{1,2}-[A-Za-z]{3}-[0-9]{4})",
        r"Income Tax Return electronically transmitted on\s+([0-9]{1,2}-[A-Za-z]{3}-[0-9]{4})",
    ]
    for pat in date_patterns:
        m = re.search(pat, text_norm, re.I)
        if m:
            filing_date = m.group(1).strip()
            break

    # Name
    name = ""
    name_patterns = [
        r"\bName\s+([A-Z][A-Z\s\.&'-]+?)\s+(?:Address|Status|Form Number|Filed u/s)\b",
        r"\bName\s+([A-Z][A-Z\s\.&'-]+)\b",
        r"\bI,\s*([A-Z][A-Z\s\.&'-]+?)\s+son\/ daughter of\b",
    ]
    for pat in name_patterns:
        m = re.search(pat, text_norm, re.I | re.S)
        if m:
            cand = re.sub(r"\s+", " ", m.group(1)).strip(" ,")
            if len(cand.split()) >= 2:
                name = cand
                break

    # Address
    address = ""
    m = re.search(r"\bAddress\s+(.*?)\s+(?:Status|Form Number|Filed u/s)\b", text_norm, re.I | re.S)
    if m:
        address = re.sub(r"\s+", " ", m.group(1)).strip(" ,")

    # Total Income
    total_income = ""
    income_patterns = [
        r"Total Income\s+[0-9A-Z]{1,3}\s+([0-9,]+(?:\.\d{2})?)",
        r"Total Income\s+([0-9,]+(?:\.\d{2})?)",
    ]
    for pat in income_patterns:
        m = re.search(pat, text_norm, re.I)
        if m:
            total_income = m.group(1).strip()
            break

    return {
        "PAN": pan,
        "AY": ay,
        "ACK": ack,
        "NAME": name,
        "DATE": filing_date,
        "INCOME": total_income,
        "ADDRESS": address,
    }

# ---------------- FORM 26AS PARSER ---------------- #

def parse_26as_details(text):
    text_norm = text.replace("\xa0", " ")
    lines = [re.sub(r"\s+", " ", l).strip() for l in text_norm.splitlines() if l.strip()]

    # Basic details
    pan = ""
    ay = ""
    name = ""

    m = re.search(r"Permanent Account Number \(PAN\)\s+([A-Z]{5}[0-9]{4}[A-Z])", text_norm)
    if m:
        pan = m.group(1).strip()

    m = re.search(r"Assessment Year\s+(\d{4}-\d{2})", text_norm)
    if m:
        ay = m.group(1).strip()

    m = re.search(
        r"Name of Assessee\s+([A-Z][A-Z\s\.\&'-]+?)(?=\s+Address of Assessee|\s+Address\b|\n|$)",
        text_norm,
        re.S
    )
    if m:
        name = re.sub(r"\s+", " ", m.group(1)).strip(" ,")
        name = re.sub(r"\s+[A-Z]$", "", name).strip()

    # Parse Part I deductor blocks
    deductor_blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        m = re.match(
            r"^\d+\s+(.+?)\s+([A-Z]{4}[A-Z0-9]{6})\s+(-?[0-9,]+\.\d{2})\s+(-?[0-9,]+\.\d{2})\s+(-?[0-9,]+\.\d{2})$",
            line
        )

        if m:
            deductor_name = re.sub(r"\s+", " ", m.group(1)).strip(" ,")
            tan = m.group(2).strip()
            total_paid = m.group(3).replace(",", "").strip()

            detail_lines = []
            j = i + 1
            while j < len(lines):
                nxt = lines[j]

                next_summary = re.match(
                    r"^\d+\s+(.+?)\s+([A-Z]{4}[A-Z0-9]{6})\s+(-?[0-9,]+\.\d{2})\s+(-?[0-9,]+\.\d{2})\s+(-?[0-9,]+\.\d{2})$",
                    nxt
                )
                if next_summary:
                    break
                if nxt.startswith("PART-II") or nxt.startswith("PART-II-") or nxt.startswith("PART-III") or nxt.startswith("PART III"):
                    break

                detail_lines.append(nxt)
                j += 1

            has_192 = any(re.search(r"^\d+\s+192\b", dl) for dl in detail_lines)

            deductor_blocks.append({
                "name": deductor_name,
                "tan": tan,
                "total_paid": total_paid,
                "has_192": has_192,
            })

            i = j
            continue

        i += 1

    deductors_192 = [d for d in deductor_blocks if d["has_192"]]

    if not deductors_192:
        return {
            "PAN": pan,
            "AY": ay,
            "NAME": name,
            "HAS_192": False,
            "DEDUCTOR": "NIL",
            "AMOUNT": "NIL"
        }

    deductor_names = ",".join([d["name"] for d in deductors_192])

    total_sum = 0.0
    for d in deductors_192:
        try:
            total_sum += float(d["total_paid"])
        except:
            pass

    amount_str = f"{total_sum:.2f}" if total_sum > 0 else "NIL"

    return {
        "PAN": pan,
        "AY": ay,
        "NAME": name,
        "HAS_192": True,
        "DEDUCTOR": deductor_names if deductor_names else "NIL",
        "AMOUNT": amount_str
    }

# ---------------- COMMON HELPERS ---------------- #

def clean_filename_component(value):
    if value is None:
        return ""
    value = str(value).strip()
    value = re.sub(r"\s+", " ", value)
    value = value.replace("/", "_").replace("\\", "_")
    value = value.replace(":", "-").replace("*", "")
    value = value.replace("?", "").replace('"', "")
    value = value.replace("<", "").replace(">", "").replace("|", "")
    return value.strip(" .")

def extract_pdf_text(path):
    reader = PdfReader(path)
    text_parts = []
    for p in reader.pages:
        try:
            page_text = p.extract_text() or ""
            text_parts.append(page_text)
        except:
            pass
    return "\n".join(text_parts)

def copy_preserve_relpath(src_file, src_root, dst_root, new_basename=None):
    rel_dir = os.path.relpath(os.path.dirname(src_file), src_root)
    out_dir = os.path.join(dst_root, rel_dir) if rel_dir != "." else dst_root
    os.makedirs(out_dir, exist_ok=True)

    if new_basename:
        dst = os.path.join(out_dir, new_basename)
    else:
        dst = os.path.join(out_dir, os.path.basename(src_file))

    shutil.copy2(src_file, dst)
    return dst

def zip_folder(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(folder_path):
            for f in files:
                full_path = os.path.join(root, f)
                arcname = os.path.relpath(full_path, folder_path)
                z.write(full_path, arcname)

def save_excel_bytes(df):
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)
    bio.seek(0)
    return bio.getvalue()

# ---------------- STREAMLIT ---------------- #

st.set_page_config(layout="centered")
st.title("📄 Document Generator")

tab1, tab2, tab3, tab4 = st.tabs([
    "🖼 Images Report",
    "📑 PDF Combiner",
    "📂 ITR File Rename",
    "📂 FORM 26AS File Rename"
])

if "itr_zip_bytes" not in st.session_state:
    st.session_state["itr_zip_bytes"] = None
if "itr_excel_bytes" not in st.session_state:
    st.session_state["itr_excel_bytes"] = None
if "itr_ready" not in st.session_state:
    st.session_state["itr_ready"] = False

if "f26_zip_bytes" not in st.session_state:
    st.session_state["f26_zip_bytes"] = None
if "f26_excel_bytes" not in st.session_state:
    st.session_state["f26_excel_bytes"] = None
if "f26_ready" not in st.session_state:
    st.session_state["f26_ready"] = False

# ================= TAB 1 ================= #

with tab1:

    st.markdown("**Output filename (without extension)**")
    filename = st.text_input("", "visit_report")

    st.markdown("**Output Format**")
    output_format = st.radio("",["PDF","Word","Both"])


    visit_uploads = {}

    for visit in VISIT_TYPES:
        st.markdown(f"**Upload images for {visit}**")
        visit_uploads[visit] = st.file_uploader(
            "",
            type=["jpg","jpeg","png"],
            accept_multiple_files=True,
            key=visit
        )

    if st.button("Generate Images Report"):

        temp_dir = tempfile.mkdtemp()
        visit_images={}
        total=0

        for visit,files in visit_uploads.items():

            paths=[]

            if files:
                for file in files:
                    safe = safe_name(visit)
                    path = os.path.join(temp_dir,f"{safe}_{file.name}")
                    with open(path,"wb") as f:
                        f.write(file.getbuffer())
                    paths.append(path)

            visit_images[visit]=paths
            total+=len(paths)

        if total==0:
            st.error("Upload at least one image")
            st.stop()

        pdf_path=os.path.join(temp_dir,f"{filename}.pdf")
        word_path=os.path.join(temp_dir,f"{filename}.docx")

        if output_format in ["PDF","Both"]:
            generate_pdf(visit_images,pdf_path)

        if output_format in ["Word","Both"]:
            generate_word(visit_images,word_path)

        st.success("Generated ✅")

        if output_format in ["PDF","Both"]:
            st.download_button("⬇ Download PDF",open(pdf_path,"rb"),file_name=f"{filename}.pdf")

        if output_format in ["Word","Both"]:
            st.download_button("⬇ Download Word",open(word_path,"rb"),file_name=f"{filename}.docx")

# ================= TAB 2 ================= #

with tab2:

    st.markdown("**Combined PDF filename**")
    pdf_filename = st.text_input("", "combined_documents")


    uploads = {}

    for cat in PDF_CATEGORIES:
        st.markdown(f"**{cat}**")
        uploads[cat] = st.file_uploader(
            "",
            type=["pdf"],
            accept_multiple_files=True,
            key=f"pdf_{cat}"
        )

    if st.button("Generate Combined PDF"):

        temp_dir=tempfile.mkdtemp()
        category_files={}
        total=0

        for cat,files in uploads.items():

            paths=[]

            if files:
                for f in files:
                    path=os.path.join(temp_dir,f"{safe_name(cat)}_{f.name}")
                    with open(path,"wb") as out:
                        out.write(f.getbuffer())
                    paths.append(path)

            category_files[cat]=paths
            total+=len(paths)

        if total==0:
            st.error("Upload at least one PDF")
            st.stop()

        output=os.path.join(temp_dir,f"{pdf_filename}.pdf")

        merge_pdfs(category_files,output)

        st.success("Combined PDF Ready ✅")

        st.download_button("⬇ Download Combined PDF",open(output,"rb"),file_name=f"{pdf_filename}.pdf")


# ================= TAB 3 ================= #

with tab3:

    st.markdown("**Upload ZIP containing ITR PDFs**")
    itr_zip = st.file_uploader("", type=["zip"], key="itr_zip")

    if st.button("Process ITR Files"):

        if not itr_zip:
            st.error("Upload ZIP file")
            st.stop()

        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, "input_itr.zip")

        with open(zip_path, "wb") as f:
            f.write(itr_zip.getbuffer())

        extract_dir = os.path.join(temp_dir, "extract")
        os.makedirs(extract_dir, exist_ok=True)

        with zipfile.ZipFile(zip_path, "r") as z:
            z.extractall(extract_dir)

        output_dir = os.path.join(temp_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        records = []

        for root, _, files in os.walk(extract_dir):
            for file in files:
                src = os.path.join(root, file)

                if file.lower().endswith(".pdf"):
                    text = extract_pdf_text(src)
                    d = parse_itr_details(text)

                    newname = (
                        f"{clean_filename_component(d['PAN'])}_"
                        f"{clean_filename_component(d['AY'])}_"
                        f"{clean_filename_component(d['ACK'])}_"
                        f"{clean_filename_component(d['NAME'])}_"
                        f"{clean_filename_component(d['DATE'])}_"
                        f"{clean_filename_component(d['INCOME'])}_ITR.pdf"
                    )

                    copy_preserve_relpath(src, extract_dir, output_dir, newname)

                    records.append({
                        "PAN NO.": d["PAN"],
                        "ASSESSMENT YEAR": d["AY"],
                        "ACKNOWLEDGEMENT NO.": d["ACK"],
                        "NAME": d["NAME"],
                        "DATE OF FILING": d["DATE"],
                        "TOTAL INCOME": d["INCOME"],
                        "ADDRESS": d["ADDRESS"],
                    })
                else:
                    copy_preserve_relpath(src, extract_dir, output_dir)

        df = pd.DataFrame(records)

        excel_bytes = save_excel_bytes(df)

        zip_out = os.path.join(temp_dir, "itr_output.zip")
        zip_folder(output_dir, zip_out)

        with open(zip_out, "rb") as f:
            zip_bytes = f.read()

        st.session_state["itr_zip_bytes"] = zip_bytes
        st.session_state["itr_excel_bytes"] = excel_bytes
        st.session_state["itr_ready"] = True

        st.success("Processing Complete ✅")

    if st.session_state["itr_ready"]:
        st.download_button(
            "⬇ Download Renamed ZIP",
            data=st.session_state["itr_zip_bytes"],
            file_name="itr_files.zip",
            mime="application/zip",
            key="itr_zip_download"
        )

        st.download_button(
            "⬇ Download Excel",
            data=st.session_state["itr_excel_bytes"],
            file_name="itr_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="itr_excel_download"
        )

# ================= TAB 4 ================= #

with tab4:

    st.markdown("**Upload ZIP containing FORM 26AS PDFs**")
    f26_zip = st.file_uploader("", type=["zip"], key="f26_zip")

    if st.button("Process FORM 26AS Files"):

        if not f26_zip:
            st.error("Upload ZIP file")
            st.stop()

        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, "input_26as.zip")

        with open(zip_path, "wb") as f:
            f.write(f26_zip.getbuffer())

        extract_dir = os.path.join(temp_dir, "extract")
        os.makedirs(extract_dir, exist_ok=True)

        with zipfile.ZipFile(zip_path, "r") as z:
            z.extractall(extract_dir)

        output_dir = os.path.join(temp_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        records = []

        for root, _, files in os.walk(extract_dir):
            for file in files:
                src = os.path.join(root, file)

                if file.lower().endswith(".pdf"):
                    text = extract_pdf_text(src)
                    d = parse_26as_details(text)

                    newname = (
                        f"{clean_filename_component(d['PAN'])}_"
                        f"{clean_filename_component(d['AY'])}_"
                        f"{clean_filename_component(d['NAME'])}_"
                        f"192_"
                        f"{clean_filename_component(d['DEDUCTOR'])}_"
                        f"{clean_filename_component(d['AMOUNT'])}_"
                        f"FORM26AS.pdf"
                    )

                    copy_preserve_relpath(src, extract_dir, output_dir, newname)

                    records.append({
                        "PAN NO.": d["PAN"],
                        "ASSESSMENT YEAR": d["AY"],
                        "NAME OF ASSESSEE": d["NAME"],
                        "192": bool(d["HAS_192"]),
                        "NAME OF DEDUCTER": d["DEDUCTOR"] if d["HAS_192"] else "NIL",
                        "TOTAL AMOUNT PAID/CREDITED": d["AMOUNT"] if d["HAS_192"] else "NIL",
                    })
                else:
                    copy_preserve_relpath(src, extract_dir, output_dir)

        df = pd.DataFrame(records)

        excel_bytes = save_excel_bytes(df)

        zip_out = os.path.join(temp_dir, "form26as_output.zip")
        zip_folder(output_dir, zip_out)

        with open(zip_out, "rb") as f:
            zip_bytes = f.read()

        st.session_state["f26_zip_bytes"] = zip_bytes
        st.session_state["f26_excel_bytes"] = excel_bytes
        st.session_state["f26_ready"] = True

        st.success("Processing Complete ✅")

    if st.session_state["f26_ready"]:
        st.download_button(
            "⬇ Download Renamed ZIP",
            data=st.session_state["f26_zip_bytes"],
            file_name="form26as_files.zip",
            mime="application/zip",
            key="f26_zip_download"
        )

        st.download_button(
            "⬇ Download Excel",
            data=st.session_state["f26_excel_bytes"],
            file_name="form26as_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="f26_excel_download"
        )